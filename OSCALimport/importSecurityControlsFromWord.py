import docx
from SystemSecurityPlans import models
import logging


def changeRoll(old_role,new_role):
    controls = models.system_control.objects.filter(control_responsible_roles=models.user_role.objects.filter(title=old_role)[0].pk)
    for item in controls:
        item.control_responsible_roles.add(models.user_role.objects.filter(title=new_role)[0].pk)
        item.control_responsible_roles.remove(models.user_role.objects.filter(title=old_role)[0].pk)
    try:
        models.user_role.objects.filter(title=old_role)[0].delete()
    except:
        print('role does not exist')


def delUnusedRoles():
    r = models.user_role.objects.all()
    for item in r:
        if item.system_control_set.count() == 0:
            print('deleting ' + item.title)
            item.delete()


def cleanUpRolesAfterImport():
    changeRoll('Information system', 'Information System')
    changeRoll('CIO', 'System Owner')
    changeRoll('Technology Services', 'System Administrators')
    changeRoll('IA', 'Information Assurance Team')
    changeRoll('Technology Services Team Lead', 'Operations Team Lead')
    changeRoll('Administrators', 'System Administrators')
    changeRoll('Information Assurance Team (IAT)', 'Information Assurance Team')
    changeRoll('security engineers', 'Information Assurance Team')
    changeRoll('All', 'organization')
    changeRoll('Organization’s designated personnel or management', 'organization')
    changeRoll('System owners', 'system owner')
    changeRoll('developers', 'Developers')
    changeRoll('MAX.gov systems', 'Information System')

def listRolesWithControlCount():
    r = models.user_role.objects.all()
    role_dictionary = {}
    for role in r:
        role_dictionary[role.title] = role.system_control_set.count()
    sort_roles = sorted(role_dictionary.items(), key=lambda x: x[1], reverse=True)

    for i in sort_roles:
        print(i[0], i[1])

def listDataTables(d):
    for para in d.paragraphs:
        if para.text[0:5] == 'Table':
            print(para.text)

def listTables(document):
    d = docx.Document(document)
    for table in d.tables:
        row1 = []
        for cell in table.row_cells(0):
            row1.append(cell.text)
        print('|' + '|'.join(row1))

def getCheckedOptions(cell):
    results = []
    cell_elm = cell._element
    # checkBoxes = cell_elm.xpath('.//w14:checkBox')
    checkBoxes = cell_elm.xpath(".//*[local-name()='checked']")
    labels = cell_elm.xpath(".//*[local-name()='t']")
    c = 0
    for item in labels:
        if type(item.text) == str and len(item.text) > 1:
            if item.text != "Implementation Status (check all that apply):" \
                    and item.text != "Control Origination (check all that apply):" \
                    and c < len(checkBoxes):
                if checkBoxes[c].values()[0] == '1':
                    results.append(item.text)
                c = c + 1
    return ','.join(results)

def cleanData():
    models.system_control.objects.all().delete()
    models.property.objects.all().delete()
    models.control_parameter.objects.all().delete()
    models.control_statement.objects.all().delete()
    models.user_role.objects.all().delete()

    from django.db import connection

    reset_countersSQL = "delete from sqlite_sequence where name in ('system_control','property','control_parameter','control_statement','user_role');"

    with connection.cursor() as cursor:
        cursor.execute(reset_countersSQL)



def main(document):
    logFile = '/home/parallels/PycharmProjects/OSCALweb/importSecurityControlsFromWord.log'
    # open(logFile,'w').close()
    logging.basicConfig(#filename=logFile,
                        filemode='w',
                        format='%(name)s - %(levelname)s - %(message)s',
                        level=logging.DEBUG
                        )
    d = docx.Document(document)
    logging.debug('Starting import from ' + document)
    for table in d.tables:
        firstRow = True
        # There are 2 kinds of tables, Control Summary and solution
        # Sometimes the solution table has only 1 column and sometimes it has 2
        columns = len(table.rows[0].cells)
        if table.rows[0].cells[0].text != table.rows[0].cells[columns-1].text:
            if table.rows[0].cells[1].text == "Control Summary Information" or table.rows[0].cells[1].text == "Control Enhancement Summary Information":
                logging.debug('Starting import for ' + table.rows[0].cells[0].text + ' control')
                newControl = models.system_control(control_id=table.rows[0].cells[0].text)
                newControl.save()
                logging.debug('Control ' + table.rows[0].cells[0].text + ' created')
                # Find and Add user_role
                roleText = table.rows[1].cells[0].text
                roleList = (roleText[roleText.find(':') + 1:]).split(',')
                for item in roleList:
                    if item.find('and') > 0:
                        sub_items = item.split('and')
                        for sub_item in sub_items:
                            roleList.append(sub_item)
                        break
                    item = item.strip()
                    logging.debug('Adding role ' + item)
                    newRole, created = models.user_role.objects.get_or_create(title=item[:100], shortName=item[:25], desc=item[:100])
                    if created: newRole.save()
                    newControl.control_responsible_roles.add(newRole.id)
                # Find and Add Parameters, Implementation Status, and Control Origination
                for row in table.rows:
                    content = row.cells[0].text
                    if content[0:9] == "Parameter":
                        logging.debug('Adding parameter ' + content)
                        newParameter, created = models.control_parameter.objects.get_or_create(control_parameter_id=content[10:content.find(':')].strip()[:25], value=content[content.find(':') + 1:].strip())
                        if created: newParameter.save()
                        newControl.control_parameters.add(newParameter.id)
                    elif content[0:14] == "Implementation":
                        logging.debug('*****DEBUG*****' + content)
                        content = getCheckedOptions(row.cells[0])
                        logging.debug('Adding Implementation Status ' + content)
                        newControl.control_status = content
                    elif content[0:7] == "Control":
                        content = getCheckedOptions(row.cells[0])
                        logging.debug('Adding Control Origination ' + content)
                        newControl.control_origination = content
            else:
                msg = '|'
                for item in table.rows[0].cells:
                    msg = msg + item.text + '|'
                logging.debug('Table not imported. First Row text was ' + msg)
        else:
            rowCount = 0
            conID = ''
            for row in table.rows:
                #skip header row
                if firstRow:
                    firstRow = False
                    t = row.cells[0].text
                    conID = t[0:t.find(' What')]
                else:
                    logging.debug('Importing control_statement for ' + conID)
                    control = models.system_control.objects.get(control_id=conID)
                    # Handle the case where the table has only 1 column
                    if len(row.cells) == 1:
                        controlName = 'Part a'
                        controlValue = row.cells[0].text
                    else:
                        controlName = row.cells[0].text
                        controlValue = row.cells[1].text
                    logging.debug('Adding statement ' + controlName + ': ' + controlValue)
                    control.control_statements.create(control_statement_id=controlName, control_statement_text=controlValue)
                    rowCount = rowCount+1


import docx
from SystemSecurityPlans import models
import logging
from OSCALimport.importSecurityControlsFromWord import getCheckedOptions, main, cleanData, listRolesWithControlCount, changeRoll, delUnusedRoles
f = 'MAX.gov System Security Plan Controls.docx'



